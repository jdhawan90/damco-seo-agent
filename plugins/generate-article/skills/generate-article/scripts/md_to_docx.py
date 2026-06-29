#!/usr/bin/env python
"""
md_to_docx.py  --  Convert a constrained-Markdown article into the 3-part .docx
deliverable used by the Damco content pipeline.

It produces, in one document:
  1. An "SEO Metadata" block (meta title, meta description, keywords, target URL).
  2. The full article with real Word heading styles (Title / Heading 2 / Heading 3),
     bullet & numbered lists, comparison tables, callout boxes, and inline
     hyperlinks on every statistic.
  3. An auto-computed Keyword Frequency Table (counts are measured from the final
     body text, never estimated) wherever {{KEYWORD_FREQUENCY_TABLE}} appears.

USAGE:
    python md_to_docx.py <input.md> <output.docx>

MARKDOWN CONTRACT
-----------------
A metadata block at the very top, between <!--META and -->:

    <!--META
    title: Role of AI in Healthcare App Development
    meta_title: Role of AI in Healthcare App Development | Damco
    meta_description: 150-160 char description here.
    platform: SEO Articles
    primary_keywords: ai in healthcare app development, ai healthcare apps
    secondary_keywords: healthcare app development company, healthcare software development services
    cta_url: https://www.damcogroup.com/healthcare/healthcare-app-development
    brief: Why this title/angle was chosen (the team's one-line direction).
    -->

`primary_keywords` and `secondary_keywords` are comma-separated and may each list
several phrases; the keyword frequency table is computed for every one of them.

Body syntax supported:
    ## H2 / ### H3 / #### H4          -> Word heading styles
    plain lines                       -> paragraphs (blank line separates)
    - bullet  /  * bullet             -> bulleted list
    1. item                           -> numbered list
    > callout line                    -> shaded callout box (e.g. Key Takeaways)
    | a | b |  + |---|---| separator   -> table
    [text](https://url)               -> inline hyperlink
    **bold**                          -> bold run
    {{KEYWORD_FREQUENCY_TABLE}}        -> auto-generated keyword table
"""
import sys
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = "1F3864"
TEAL = "2A9D8F"
LIGHT = "EAF3F1"   # light teal fill for callouts
META_FILL = "F2F2F2"

INLINE = re.compile(r'\[([^\]]+)\]\((https?://[^)\s]+)\)|\*\*([^*]+)\*\*')


# --------------------------------------------------------------------------- #
# Low-level helpers
# --------------------------------------------------------------------------- #
def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color'); color.set(qn('w:val'), NAVY); rPr.append(color)
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
    run.append(rPr)
    t = OxmlElement('w:t'); t.set(qn('xml:space'), 'preserve'); t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_runs(paragraph, text):
    """Render inline markdown (links + bold) into a paragraph."""
    pos = 0
    for m in INLINE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        if m.group(1) is not None:            # hyperlink
            add_hyperlink(paragraph, m.group(1), m.group(2))
        elif m.group(3) is not None:          # bold
            paragraph.add_run(m.group(3)).bold = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


# --------------------------------------------------------------------------- #
# Parsing
# --------------------------------------------------------------------------- #
def parse_meta(text):
    meta = {}
    m = re.search(r'<!--META(.*?)-->', text, re.DOTALL)
    body = text
    if m:
        for line in m.group(1).strip().splitlines():
            if ':' in line:
                k, v = line.split(':', 1)
                meta[k.strip().lower()] = v.strip()
        body = text[m.end():]
    return meta, body.strip()


def keyword_count(body, phrase):
    """Count whole-phrase, case-insensitive occurrences in plain text."""
    # strip markdown link targets so URLs don't inflate counts
    plain = re.sub(r'\]\((https?://[^)\s]+)\)', '] ', body)
    plain = re.sub(r'[#*>|`]', ' ', plain)
    pattern = r'\b' + re.escape(phrase.strip().lower()) + r'\b'
    return len(re.findall(pattern, plain.lower()))


# --------------------------------------------------------------------------- #
# Builders
# --------------------------------------------------------------------------- #
def build_meta_block(doc, meta):
    body_wordcount = meta.get('_wordcount', '')
    rows = [
        ("Meta Title", meta.get('meta_title', meta.get('title', ''))),
        ("Meta Description", meta.get('meta_description', '')),
        ("Target Platform", meta.get('platform', '')),
        ("Primary Keyword(s)", meta.get('primary_keywords', '')),
        ("Secondary Keyword(s)", meta.get('secondary_keywords', '')),
        ("Brand CTA / Target URL", meta.get('cta_url', '')),
        ("Word Count (body)", str(body_wordcount)),
    ]
    if meta.get('brief'):
        rows.insert(3, ("Editorial Direction / Brief", meta.get('brief', '')))
    heading = doc.add_paragraph()
    run = heading.add_run("SEO Metadata (for editorial / publishing — not part of article body)")
    run.bold = True; run.font.size = Pt(11); run.font.color.rgb = RGBColor.from_string(NAVY)

    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = ''
        cells[0].paragraphs[0].add_run(label).bold = True
        shade_cell(cells[0], META_FILL)
        cells[1].text = ''
        add_runs(cells[1].paragraphs[0], value)
        # flag a missing/short meta description
        if label == "Meta Description" and (len(value) < 120 or len(value) > 160):
            note = cells[1].add_paragraph()
            r = note.add_run("  [check length: aim 150-160 chars]")
            r.italic = True; r.font.size = Pt(8); r.font.color.rgb = RGBColor.from_string("C00000")
    doc.add_paragraph()


def build_keyword_table(doc, meta, body):
    doc.add_paragraph()  # spacing
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, label in enumerate(("Keyword", "Type", "Count in Article")):
        hdr[i].text = ''
        run = hdr[i].paragraphs[0].add_run(label); run.bold = True
        shade_cell(hdr[i], LIGHT)

    def add_kw_rows(kw_string, kind):
        for kw in [k.strip() for k in kw_string.split(',') if k.strip()]:
            count = keyword_count(body, kw)
            cells = table.add_row().cells
            cells[0].text = kw
            cells[1].text = kind
            cells[2].text = str(count)
            if count == 0:
                shade_cell(cells[2], "FCE4E4")  # red-ish flag for unused keyword

    add_kw_rows(meta.get('primary_keywords', ''), 'Primary')
    add_kw_rows(meta.get('secondary_keywords', ''), 'Secondary')


def build_callout(doc, lines):
    """Render consecutive '>' lines as a shaded single-cell box."""
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.rows[0].cells[0]
    shade_cell(cell, LIGHT)
    cell.text = ''
    first = True
    for ln in lines:
        content = ln[1:].strip() if ln.startswith('>') else ln.strip()
        if not content:
            continue
        if content.startswith('- ') or content.startswith('* '):
            p = cell.add_paragraph(style='List Bullet')
            add_runs(p, content[2:].strip())
        else:
            p = cell.paragraphs[0] if first and not cell.paragraphs[0].text else cell.add_paragraph()
            # bold a "Label:" lead-in
            add_runs(p, content)
        first = False
    doc.add_paragraph()


def build_table(doc, rows):
    # rows: list of lists of cell strings; first row is header, second is separator (dropped)
    data = [r for i, r in enumerate(rows) if not (i == 1 and all(set(c) <= set('-: ') for c in r))]
    if not data:
        return
    ncol = max(len(r) for r in data)
    table = doc.add_table(rows=0, cols=ncol)
    table.style = 'Table Grid'
    for ri, row in enumerate(data):
        cells = table.add_row().cells
        for ci in range(ncol):
            text = row[ci] if ci < len(row) else ''
            cells[ci].text = ''
            add_runs(cells[ci].paragraphs[0], text)
            if ri == 0:
                for run in cells[ci].paragraphs[0].runs:
                    run.bold = True
                shade_cell(cells[ci], LIGHT)
    doc.add_paragraph()


def split_table_row(line):
    line = line.strip()
    if line.startswith('|'):
        line = line[1:]
    if line.endswith('|'):
        line = line[:-1]
    return [c.strip() for c in line.split('|')]


# --------------------------------------------------------------------------- #
# Main render loop
# --------------------------------------------------------------------------- #
def render(meta, body, out_path):
    doc = Document()
    # base font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # word count of body (strip markup)
    plain = re.sub(r'\[([^\]]+)\]\((https?://[^)\s]+)\)', r'\1', body)
    plain = re.sub(r'[#*>|`]', ' ', plain)
    plain = plain.replace('{{KEYWORD_FREQUENCY_TABLE}}', ' ')
    meta['_wordcount'] = len(plain.split())

    build_meta_block(doc, meta)

    # Title
    title = doc.add_heading(meta.get('title', ''), level=0)

    lines = body.splitlines()
    i = 0
    para_buf = []

    def flush_para():
        nonlocal para_buf
        if para_buf:
            text = ' '.join(l.strip() for l in para_buf).strip()
            if text:
                p = doc.add_paragraph()
                add_runs(p, text)
            para_buf = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            flush_para()
            i += 1
            continue

        # keyword table placeholder
        if stripped == '{{KEYWORD_FREQUENCY_TABLE}}':
            flush_para()
            build_keyword_table(doc, meta, body)
            i += 1
            continue

        # headings
        hm = re.match(r'^(#{2,4})\s+(.*)$', stripped)
        if hm:
            flush_para()
            level = len(hm.group(1))
            doc.add_heading(hm.group(2).strip(), level=level)
            i += 1
            continue

        # callout box
        if stripped.startswith('>'):
            flush_para()
            block = []
            while i < len(lines) and lines[i].strip().startswith('>'):
                block.append(lines[i].strip())
                i += 1
            build_callout(doc, block)
            continue

        # tables
        if stripped.startswith('|'):
            flush_para()
            block = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                block.append(split_table_row(lines[i]))
                i += 1
            build_table(doc, block)
            continue

        # unordered list
        if re.match(r'^[-*]\s+', stripped):
            flush_para()
            while i < len(lines) and re.match(r'^[-*]\s+', lines[i].strip()):
                p = doc.add_paragraph(style='List Bullet')
                add_runs(p, re.sub(r'^[-*]\s+', '', lines[i].strip()))
                i += 1
            continue

        # ordered list
        if re.match(r'^\d+\.\s+', stripped):
            flush_para()
            while i < len(lines) and re.match(r'^\d+\.\s+', lines[i].strip()):
                p = doc.add_paragraph(style='List Number')
                add_runs(p, re.sub(r'^\d+\.\s+', '', lines[i].strip()))
                i += 1
            continue

        # default: paragraph text
        para_buf.append(line)
        i += 1

    flush_para()
    doc.save(out_path)
    print(f"Wrote {out_path}  (body word count: {meta['_wordcount']})")


def main():
    if len(sys.argv) != 3:
        print(__doc__)
        sys.exit(1)
    with open(sys.argv[1], encoding='utf-8') as f:
        text = f.read()
    meta, body = parse_meta(text)
    render(meta, body, sys.argv[2])


if __name__ == '__main__':
    main()
